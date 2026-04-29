from docx import Document
from docx.oxml import OxmlElement
from docx.shared import RGBColor
from docx.text.paragraph import Paragraph

from core.parser.docx_parser import iter_docx_paragraphs


def _insert_paragraph_after(paragraph, text):
    new_paragraph_element = OxmlElement("w:p")
    paragraph._p.addnext(new_paragraph_element)
    new_paragraph = Paragraph(new_paragraph_element, paragraph._parent)

    label = new_paragraph.add_run("Assessor feedback: ")
    label.bold = True
    label.font.color.rgb = RGBColor(31, 78, 121)

    feedback = new_paragraph.add_run(text)
    feedback.italic = True
    feedback.font.color.rgb = RGBColor(31, 78, 121)
    return new_paragraph


def annotate_docx(input_path, output_path, evaluation_results):
    doc = Document(input_path)

    paragraphs = list(iter_docx_paragraphs(doc))
    inserted_for = set()

    for q_id, result in evaluation_results.items():
        feedback = result.get("feedback", "").strip()
        anchor_index = result.get("anchor_index")
        if not feedback or anchor_index is None or anchor_index >= len(paragraphs):
            continue
        _insert_paragraph_after(paragraphs[anchor_index], feedback)
        inserted_for.add(q_id)

    missing_feedback = [
        f"{q_id}: {result.get('feedback', '').strip()}"
        for q_id, result in evaluation_results.items()
        if q_id not in inserted_for and result.get("feedback", "").strip()
    ]
    if missing_feedback:
        doc.add_paragraph()
        summary = doc.add_paragraph()
        summary.add_run("Assessor feedback for sections not located in the document: ").bold = True
        summary.add_run(" | ".join(missing_feedback))

    doc.save(output_path)
